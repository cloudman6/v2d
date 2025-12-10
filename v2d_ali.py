import os
import numpy as np
from moviepy.editor import VideoFileClip
import cv2
from skimage.metrics import structural_similarity as ssim
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import timedelta
import json
import time
import argparse
import glob
from aliyunsdkcore.acs_exception.exceptions import ClientException
from aliyunsdkcore.acs_exception.exceptions import ServerException
from aliyunsdkcore.client import AcsClient
from aliyunsdkcore.request import CommonRequest
import alibabacloud_oss_v2 as oss
from alibabacloud_oss_v2.credentials import EnvironmentVariableCredentialsProvider
import oss2
from oss2.credentials import EnvironmentVariableCredentialsProvider as Oss2CredentialsProvider

# ===================== Configuration Parameters (Adjustable) =====================
VIDEO_PATH = "input_video.mp4"  # Input video path
OUTPUT_WORD = "output.docx"  # Output Word document path
TEMP_DIR = "temp_files"  # Temporary files directory (audio, key frames)
FRAME_INTERVAL = 1  # Frame extraction interval (seconds/frame)
SSIM_THRESHOLD = 0.93  # Key frame similarity threshold (≥0.93 considered duplicate)
MERGE_SEC = 2  # Adjacent text segment merge threshold (merge if interval < 2 seconds)

# Black border cropping configuration
ENABLE_BLACK_BORDER_CROP = True  # Enable black border cropping feature
BLACK_THRESHOLD = 30  # Black border pixel threshold (0-255, smaller values are stricter)
MIN_CONTENT_HEIGHT = 100  # Minimum content height after cropping (pixels)

# Alibaba Cloud configuration
# OSS configuration (loaded from environment variables)
OSS_BUCKET = os.getenv('OSS_BUCKET', 'v2d')
OSS_REGION = os.getenv('OSS_REGION', 'cn-beijing')
OSS_ENDPOINT = os.getenv('OSS_ENDPOINT', 'oss-cn-beijing.aliyuncs.com')
OSS_ENDPOINT_INTERNAL = os.getenv('OSS_ENDPOINT_INTERNAL', 'oss-cn-beijing-internal.aliyuncs.com')
OSS_OBJECT_NAME = "temp_audio.wav"  # OSS object name
# ===================== Alibaba Cloud API Constants =====================
# Request parameters
KEY_APP_KEY = "appkey"
KEY_FILE_LINK = "file_link"
KEY_VERSION = "version"
KEY_ENABLE_WORDS = "enable_words"

# Response parameters
KEY_TASK = "Task"
KEY_TASK_ID = "TaskId"
KEY_STATUS_TEXT = "StatusText"
KEY_RESULT = "Result"

# Status values
STATUS_SUCCESS = "SUCCESS"
STATUS_RUNNING = "RUNNING"
STATUS_QUEUEING = "QUEUEING"

# ===================== Utility Functions =====================
def create_temp_dir():
    """Create temporary directory"""
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)
    return TEMP_DIR

def delete_temp_files():
    """Delete PNG images and WAV audio files in temporary directory"""
    if os.path.exists(TEMP_DIR):
        # Delete all PNG image files
        png_files = glob.glob(os.path.join(TEMP_DIR, "*.png"))
        for png_file in png_files:
            try:
                os.remove(png_file)
                # print(f"Deleted PNG file: {png_file}")
            except Exception as e:
                print(f"Failed to delete PNG file {png_file}: {e}")

        # Delete all WAV audio files
        wav_files = glob.glob(os.path.join(TEMP_DIR, "*.wav"))
        for wav_file in wav_files:
            try:
                os.remove(wav_file)
                # print(f"Deleted WAV file: {wav_file}")
            except Exception as e:
                print(f"Failed to delete WAV file {wav_file}: {e}")

        print(f"Temporary files cleanup completed: deleted {len(png_files)} PNG files and {len(wav_files)} WAV files")

def detect_black_borders(image):
    """Detect top and bottom black border areas in an image

    Args:
        image: OpenCV image object

    Returns:
        tuple: (top_crop, bottom_crop) number of pixels to crop from top and bottom
    """
    if image is None:
        return 0, 0

    height, width = image.shape[:2]

    # Convert to grayscale
    if len(image.shape) == 3:
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    else:
        gray = image

    # Calculate average brightness for each row
    row_means = []
    for y in range(height):
        row_mean = np.mean(gray[y, :])
        row_means.append(row_mean)

    # Detect top black border
    top_crop = 0
    for y in range(height):
        if row_means[y] > BLACK_THRESHOLD:
            break
        top_crop = y + 1

    # Detect bottom black border
    bottom_crop = 0
    for y in range(height - 1, -1, -1):
        if row_means[y] > BLACK_THRESHOLD:
            break
        bottom_crop = height - y

    # Ensure sufficient content area remains after cropping
    remaining_height = height - top_crop - bottom_crop
    if remaining_height < MIN_CONTENT_HEIGHT:
        # If content is too small after cropping, reduce crop amount
        excess_crop = MIN_CONTENT_HEIGHT - remaining_height
        if top_crop > bottom_crop:
            top_crop = max(0, top_crop - excess_crop // 2)
            bottom_crop = max(0, bottom_crop - excess_crop // 2)
        else:
            bottom_crop = max(0, bottom_crop - excess_crop // 2)
            top_crop = max(0, top_crop - excess_crop // 2)

    return top_crop, bottom_crop

def crop_black_borders(image_path):
    """Crop top and bottom black borders from an image

    Args:
        image_path: Image file path

    Returns:
        bool: Whether cropping was successful
    """
    try:
        # Read image
        image = cv2.imread(image_path)
        if image is None:
            print(f"Unable to read image: {image_path}")
            return False

        # Detect black borders
        top_crop, bottom_crop = detect_black_borders(image)

        # If no cropping needed
        if top_crop == 0 and bottom_crop == 0:
            print(f"No cropping needed: {image_path}")
            return True

        # Perform cropping
        height, width = image.shape[:2]
        cropped_image = image[top_crop:height-bottom_crop, :]

        # Save cropped image (overwrite original file)
        cv2.imwrite(image_path, cropped_image)
        # print(f"Cropping successful: {image_path} (crop: top{top_crop}px, bottom{bottom_crop}px)")
        return True

    except Exception as e:
        print(f"Cropping failed {image_path}: {e}")
        return False

def crop_all_keyframes(key_frames):
    """Batch crop black borders from all key frames

    Args:
        key_frames: List of key frames

    Returns:
        int: Number of successfully cropped images
    """
    if not ENABLE_BLACK_BORDER_CROP:
        print("Black border cropping feature is disabled")
        return 0

    print(f"Starting black border cropping for {len(key_frames)} key frames...")

    success_count = 0
    for i, key_frame in enumerate(key_frames):
        frame_path = key_frame["path"]
        if crop_black_borders(frame_path):
            success_count += 1

        # Output progress every 10 key frames
        if (i + 1) % 10 == 0:
            print(f"Processed {i + 1}/{len(key_frames)} key frames")

    print(f"Key frame cropping completed: {success_count}/{len(key_frames)} successful")
    return success_count

def seconds_to_hms(seconds):
    """Convert seconds to HH:MM:SS format"""
    return str(timedelta(seconds=round(seconds)))

def save_ali_result_to_cache(result_data, cache_file_path):
    """Save Alibaba Cloud API results to cache file"""
    try:
        with open(cache_file_path, 'w', encoding='utf-8') as f:
            json.dump(result_data, f, ensure_ascii=False, indent=2)
        print(f"Alibaba Cloud API results saved to cache file: {cache_file_path}")
        return True
    except Exception as e:
        print(f"Failed to save cache file: {e}")
        return False

def load_ali_result_from_cache(cache_file_path):
    """Load Alibaba Cloud API results from cache file"""
    try:
        if os.path.exists(cache_file_path):
            with open(cache_file_path, 'r', encoding='utf-8') as f:
                result_data = json.load(f)
            print(f"Loaded Alibaba Cloud API results from cache file: {cache_file_path}")
            return result_data
        else:
            print(f"Cache file does not exist: {cache_file_path}")
            return None
    except Exception as e:
        print(f"Failed to load cache file: {e}")
        return None

def video_to_audio(video_path, audio_path):
    """Extract audio from video (WAV format, 16000Hz sample rate)"""
    with VideoFileClip(video_path) as video:
        audio = video.audio
        audio.write_audiofile(audio_path, codec="pcm_s16le", fps=16000)  # WAV uncompressed format, 16000Hz sample rate, better for speech recognition
    print(f"Audio extraction completed: {audio_path} (16000Hz)")

def ali_audio_to_text(audio_file_path, cache_file_path=None, use_cache=False):
    """Use Alibaba Cloud speech recognition API for speech to text (with timestamps)

    Args:
        audio_file_path: Local audio file path
        cache_file_path: Cache file path
        use_cache: Whether to use cache mode (True: only load from cache, False: call API and save results)
    """
    # If using cache mode, try to load from cache
    if use_cache and cache_file_path:
        cached_result = load_ali_result_from_cache(cache_file_path)
        if cached_result:
            print("Using cached data, skipping API call")
            return parse_ali_result(cached_result)
        else:
            print("Cache file does not exist, cannot use cache mode")
            return []

    # Automatically upload audio file to OSS and get URL
    print("=== Starting automatic OSS upload process ===")
    print("1. Uploading audio file to OSS...")
    if not upload_audio_to_oss(audio_file_path):
        print("❌ OSS upload failed, cannot continue speech recognition")
        return []

    print("2. Generating OSS audio file URL...")
    audio_file_link = generate_oss_url()
    if not audio_file_link:
        print("❌ OSS URL generation failed, cannot continue speech recognition")
        return []

    print(f"✅ OSS audio file URL generated successfully: {audio_file_link}")
    print("=== OSS upload process completed ===")

    # Get Alibaba Cloud credentials from environment variables
    accessKeyId = os.getenv('ALIYUN_AK_ID')
    accessKeySecret = os.getenv('ALIYUN_AK_SECRET')
    appKey = os.getenv('NLS_APP_KEY')

    print(f"AccessKey ID: {accessKeyId}")
    print(f"AccessKey Secret: {accessKeySecret}")

    # Region ID, fixed value.
    REGION_ID = "cn-beijing"
    PRODUCT = "nls-filetrans"
    DOMAIN = "filetrans.cn-beijing.aliyuncs.com"
    API_VERSION = "2018-08-17"
    POST_REQUEST_ACTION = "SubmitTask"
    GET_REQUEST_ACTION = "GetTaskResult"


    # Create AcsClient instance
    client = AcsClient(accessKeyId, accessKeySecret, REGION_ID)

    # Submit audio file recognition request
    postRequest = CommonRequest()
    postRequest.set_domain(DOMAIN)
    postRequest.set_version(API_VERSION)
    postRequest.set_product(PRODUCT)
    postRequest.set_action_name(POST_REQUEST_ACTION)
    postRequest.set_method('POST')

    # New integrations should use version 4.0, existing integrations (default 2.0) can comment out this parameter if maintaining current state.
    # Set whether to output word information, default is false, requires version 4.0 when enabled.
    task = {KEY_APP_KEY: appKey, KEY_FILE_LINK: audio_file_link, KEY_VERSION: "4.0", KEY_ENABLE_WORDS: False}
    task = json.dumps(task)
    print(task)
    postRequest.add_body_params(KEY_TASK, task)

    taskId = ""
    try:
        postResponse = client.do_action_with_exception(postRequest)
        postResponse = json.loads(postResponse)
        print(postResponse)
        statusText = postResponse[KEY_STATUS_TEXT]
        if statusText == STATUS_SUCCESS:
            print("Audio file recognition request successful!")
            taskId = postResponse[KEY_TASK_ID]
        else:
            print("Audio file recognition request failed!")
            return []
    except ServerException as e:
        print(e)
        return []
    except ClientException as e:
        print(e)
        return []

    # Create CommonRequest, set task ID.
    getRequest = CommonRequest()
    getRequest.set_domain(DOMAIN)
    getRequest.set_version(API_VERSION)
    getRequest.set_product(PRODUCT)
    getRequest.set_action_name(GET_REQUEST_ACTION)
    getRequest.set_method('GET')
    getRequest.add_query_param(KEY_TASK_ID, taskId)

    # Submit audio file recognition result query request
    # Poll for recognition results until server returns status "SUCCESS", "SUCCESS_WITH_NO_VALID_FRAGMENT",
    # or an error description, then stop polling.
    statusText = ""
    result_data = None

    while True:
        try:
            getResponse = client.do_action_with_exception(getRequest)
            getResponse = json.loads(getResponse)
            # print(getResponse)
            statusText = getResponse[KEY_STATUS_TEXT]
            if statusText == STATUS_RUNNING or statusText == STATUS_QUEUEING:
                # Continue polling
                time.sleep(10)
            else:
                # Exit polling
                if statusText == STATUS_SUCCESS:
                    result_data = getResponse
                break
        except ServerException as e:
            print(e)
            break
        except ClientException as e:
            print(e)
            break

    if statusText == STATUS_SUCCESS:
        print("Audio file recognition successful!")
        # If cache file path provided, save results to cache
        if cache_file_path:
            save_ali_result_to_cache(result_data, cache_file_path)
        # Parse results and convert to standard format
        result = parse_ali_result(result_data)

        # Delete OSS audio file after successful speech recognition
        print("=== Starting OSS cleanup process ===")
        print("Cleaning up OSS audio file...")
        if not delete_oss_audio():
            print("⚠️ Warning: OSS audio file cleanup failed")
        else:
            print("✅ OSS audio file cleanup successful")
        print("=== OSS cleanup process completed ===")

        return result
    else:
        print("Audio file recognition failed!")
        # Even if recognition fails, try to clean up OSS files
        print("=== Starting OSS cleanup process (recognition failed) ===")
        print("Cleaning up OSS audio file...")
        if not delete_oss_audio():
            print("⚠️ Warning: OSS audio file cleanup failed")
        else:
            print("✅ OSS audio file cleanup successful")
        print("=== OSS cleanup process completed ===")
        return []

def upload_audio_to_oss(audio_file_path):
    """Upload audio file to OSS"""
    try:
        print(f"Starting audio file upload to OSS: {audio_file_path}")

        # Load credential information from environment variables
        credentials_provider = EnvironmentVariableCredentialsProvider()

        # Load SDK default configuration and set credential provider
        cfg = oss.config.load_default()
        cfg.credentials_provider = credentials_provider
        cfg.region = OSS_REGION
        cfg.endpoint = OSS_ENDPOINT

        # Create OSS client with configured information
        client = oss.Client(cfg)

        # Execute object upload request, upload directly from file
        result = client.put_object_from_file(
            oss.PutObjectRequest(
                bucket=OSS_BUCKET,
                key=OSS_OBJECT_NAME
            ),
            audio_file_path
        )

        if result.status_code == 200:
            print(f"Audio file upload successful: {OSS_OBJECT_NAME}")
            print(f"Status code: {result.status_code}, Request ID: {result.request_id}")
            return True
        else:
            print(f"Audio file upload failed: status code {result.status_code}")
            return False

    except Exception as e:
        print(f"OSS upload failed: {e}")
        return False

def generate_oss_url():
    """Generate pre-signed URL for OSS audio file"""
    try:
        print("Generating OSS audio file URL...")

        # Get access credentials from environment variables
        auth = oss2.ProviderAuthV4(Oss2CredentialsProvider())

        # Use internal endpoint to generate URL
        endpoint = f"https://{OSS_ENDPOINT_INTERNAL}"


        # Create Bucket object
        bucket = oss2.Bucket(auth, endpoint, OSS_BUCKET, region=OSS_REGION)

        # Generate pre-signed URL for file download, valid for 600 seconds
        # Set slash_safe to True, OSS will not escape forward slashes (/) in Object full path
        url = bucket.sign_url('GET', OSS_OBJECT_NAME, 600, slash_safe=True)

        print(f"OSS audio file URL generated successfully: {OSS_OBJECT_NAME}")
        return url

    except Exception as e:
        print(f"OSS URL generation failed: {e}")
        return None

def delete_oss_audio():
    """Delete audio file from OSS"""
    try:
        print(f"Deleting OSS audio file: {OSS_OBJECT_NAME}")

        # Load credential information from environment variables
        credentials_provider = EnvironmentVariableCredentialsProvider()

        # Load SDK default configuration and set credential provider
        cfg = oss.config.load_default()
        cfg.credentials_provider = credentials_provider
        cfg.region = OSS_REGION
        cfg.endpoint = OSS_ENDPOINT
        # Set timeout: connection timeout 30s, read timeout 300s (5 minutes)
        cfg.connect_timeout = 30
        cfg.readwrite_timeout = 300

        # Create OSS client with configured information
        client = oss.Client(cfg)

        # Execute object deletion request
        result = client.delete_object(oss.DeleteObjectRequest(
            bucket=OSS_BUCKET,
            key=OSS_OBJECT_NAME,
        ))

        if result.status_code == 204:
            print(f"OSS audio file deletion successful: {OSS_OBJECT_NAME}")
            return True
        else:
            print(f"OSS audio file deletion failed: status code {result.status_code}")
            return False

    except Exception as e:
        print(f"OSS deletion failed: {e}")
        return False

def parse_ali_result(result_data):
    """Parse Alibaba Cloud speech recognition results and convert to standard format"""
    if not result_data or KEY_RESULT not in result_data:
        print("No recognition results found")
        return []

    result = result_data[KEY_RESULT]
    if "Sentences" not in result:
        print("No sentence data found")
        return []

    sentences = result["Sentences"]
    formatted_segments = []
    seen_segments = set()  # Set for deduplication

    print(f"Raw API returned data: {len(sentences)} sentences total")

    for i, sentence in enumerate(sentences):
        # Alibaba Cloud returns time in milliseconds, convert to seconds
        start_time = sentence.get("BeginTime", 0) / 1000.0
        end_time = sentence.get("EndTime", 0) / 1000.0
        text = sentence.get("Text", "")
        channel_id = sentence.get("ChannelId", 0)
        speaker_id = sentence.get("SpeakerId", "")

        # Create unique identifier for deduplication (based on time range and text content)
        segment_key = (start_time, end_time, text)

        # If this segment has already been processed, skip duplicate
        if segment_key in seen_segments:
            # print(f"Skipping duplicate segment: channel{channel_id}, speaker{speaker_id}, {seconds_to_hms(start_time)} - {seconds_to_hms(end_time)}: {text[:30]}...")
            continue

        # Add to processed set
        seen_segments.add(segment_key)

        formatted_segments.append({
            "start": start_time,
            "end": end_time,
            "start_hms": seconds_to_hms(start_time),
            "end_hms": seconds_to_hms(end_time),
            "text": text
        })

    print(f"Speech to text completed, {len(formatted_segments)} segments (after deduplication)")
    # print("Segment details:")
    # for i, seg in enumerate(formatted_segments[:10]):  # Only show first 10 segments
    #     print(f"  Segment {i+1}: {seg['start_hms']} - {seg['end_hms']} ({seg['end']-seg['start']:.1f}s): {seg['text'][:50]}...")

    return formatted_segments

def extract_key_frames(video_path):
    """Extract key frames (with deduplication)"""
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS)  # Video frame rate
    frame_interval = int(fps * FRAME_INTERVAL)  # Extract one frame every FRAME_INTERVAL seconds (converted to frame count)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

    key_frames = []  # Store key frames: [{time, path}, ...]
    last_frame_gray = None  # Previous key frame grayscale image
    frame_count = 0

    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break

        # Extract frames at intervals
        if frame_count % frame_interval != 0:
            frame_count += 1
            continue

        # Calculate current frame timestamp (seconds)
        current_time = frame_count / fps

        # Convert to grayscale (for SSIM calculation)
        current_frame_gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

        # Determine if it's a key frame (first frame is always kept)
        is_key_frame = False
        if last_frame_gray is None:
            is_key_frame = True
        else:
            # Calculate SSIM similarity (win_size must be odd, adapt to different frame sizes)
            win_size = min(7, current_frame_gray.shape[0]//2, current_frame_gray.shape[1]//2)
            if win_size % 2 == 0:
                win_size -= 1
            sim, _ = ssim(last_frame_gray, current_frame_gray, full=True, win_size=win_size)
            if sim < SSIM_THRESHOLD:
                is_key_frame = True

        if is_key_frame:
            # Save key frame
            frame_path = os.path.join(TEMP_DIR, f"key_frame_{len(key_frames)}.png")
            cv2.imwrite(frame_path, frame)
            key_frames.append({
                "time": current_time,
                "path": frame_path
            })
            last_frame_gray = current_frame_gray
            # print(f"Saved key frame: {frame_path} (time: {seconds_to_hms(current_time)})")

        frame_count += 1

    cap.release()
    print(f"Key frame extraction completed, {len(key_frames)} frames")
    return key_frames

def get_video_duration(video_path):
    """Get video total duration (seconds)"""
    with VideoFileClip(video_path) as video:
        duration = video.duration
    return duration

def match_text_and_frames(text_segments, key_frames, video_path):
    """Segment text fragments using key frame times: simplified logic"""
    # Get video total duration
    video_duration = get_video_duration(video_path)

    # Extract key frame times and sort them
    key_frame_times = sorted([kf["time"] for kf in key_frames])

    print(f"Key frame segmentation started: {len(text_segments)} text segments, {len(key_frames)} key frames")
    print(f"Video total duration: {seconds_to_hms(video_duration)}")
    # print(f"Key frame times: {key_frame_times}")

    # Create segment data: each key frame corresponds to one segment
    segments_data = []

    # If no key frames, create one segment containing all text fragments
    if not key_frame_times:
        segments_data.append({
            "segment_index": 1,
            "start_time": 0,
            "end_time": video_duration,
            "start_hms": seconds_to_hms(0),
            "end_hms": seconds_to_hms(video_duration),
            "key_frame": None,
            "text_segments": text_segments
        })
        # print(f"  Segment 1: {seconds_to_hms(0)} - {seconds_to_hms(video_duration)} -> {len(text_segments)} text segments [No key frame]")
    else:
        # First segment: from 0 to second key frame time (if second key frame exists)
        first_keyframe = key_frames[0]
        end_time = key_frame_times[1] if len(key_frame_times) > 1 else video_duration

        # Find all text fragments belonging to this time period
        segment_texts = []
        for seg in text_segments:
            if 0 <= seg["start"] < end_time:
                segment_texts.append(seg)

        segments_data.append({
            "segment_index": 1,
            "start_time": 0,
            "end_time": end_time,
            "start_hms": seconds_to_hms(0),
            "end_hms": seconds_to_hms(end_time),
            "key_frame": first_keyframe,
            "text_segments": segment_texts
        })
        # print(f"  Segment 1: {seconds_to_hms(0)} - {seconds_to_hms(end_time)} -> {len(segment_texts)} text segments [Key frame: {first_keyframe['time']:.1f}s]")

        # Middle segments: from i-th key frame to (i+1)-th key frame (i from 1 to n-2)
        for i in range(1, len(key_frame_times) - 1):
            start_time = key_frame_times[i]
            end_time = key_frame_times[i + 1]
            key_frame = key_frames[i]

            # Find all text fragments belonging to this time period
            segment_texts = []
            for seg in text_segments:
                if start_time <= seg["start"] < end_time:
                    segment_texts.append(seg)

            segments_data.append({
                "segment_index": i + 1,
                "start_time": start_time,
                "end_time": end_time,
                "start_hms": seconds_to_hms(start_time),
                "end_hms": seconds_to_hms(end_time),
                "key_frame": key_frame,
                "text_segments": segment_texts
            })
            # print(f"  Segment {i+1}: {seconds_to_hms(start_time)} - {seconds_to_hms(end_time)} -> {len(segment_texts)} text segments [Key frame: {key_frame['time']:.1f}s]")

        # Last segment: from last key frame to video end
        if len(key_frame_times) > 1:
            last_keyframe = key_frames[-1]
            start_time = key_frame_times[-1]
            end_time = video_duration

            # Find all text fragments belonging to this time period
            segment_texts = []
            for seg in text_segments:
                if start_time <= seg["start"] <= end_time:
                    segment_texts.append(seg)

            segments_data.append({
                "segment_index": len(key_frame_times),
                "start_time": start_time,
                "end_time": end_time,
                "start_hms": seconds_to_hms(start_time),
                "end_hms": seconds_to_hms(end_time),
                "key_frame": last_keyframe,
                "text_segments": segment_texts
            })
            # print(f"  Segment {len(key_frame_times)}: {seconds_to_hms(start_time)} - {seconds_to_hms(end_time)} -> {len(segment_texts)} text segments [Key frame: {last_keyframe['time']:.1f}s]")

    # Validation: ensure all text fragments are assigned to some segment
    assigned_texts = sum(len(seg["text_segments"]) for seg in segments_data)
    if assigned_texts != len(text_segments):
        raise Exception(f"Text fragment assignment error: {len(text_segments)} total fragments, only {assigned_texts} assigned")

    print(f"Key frame segmentation completed, {len(segments_data)} segments")
    return segments_data

def split_text_by_keyframes(segments_data):
    """Simplify text segmentation: directly use segment data"""
    # In the new segmentation logic, segmentation is already done, directly return segment data
    print(f"Text segmentation completed: {len(segments_data)} segments total")
    return segments_data

def generate_word(segments_data, output_path):
    """Generate Word document: simplified format - key frames + corresponding text segments"""
    doc = Document()

    print(f"Generating Word document: {len(segments_data)} segments total")

    # Set document title
    title = doc.add_heading("Video Notes", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Traverse segment data and insert content
    for segment in segments_data:
        # Debug information
        # print(f"Generating segment {segment['segment_index']}: key_frame={segment['key_frame'] is not None}")
        # if segment["key_frame"]:
        #     print(f"  Key frame time: {segment['key_frame']['time']:.1f} seconds")

        # Insert key frame title
        if segment["key_frame"]:
            # keyframe_heading = doc.add_heading(
            #     f"Key Frame {segment['segment_index']} ({segment['key_frame']['time']:.1f} seconds)",
            #     level=2
            # )

            # Insert key frame image
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(segment["key_frame"]["path"], width=Inches(6))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # else:
            # For segments without key frames (like the last segment), only show time period
            # time_heading = doc.add_heading(
            #     f"Segment {segment['segment_index']}: {segment['start_hms']} - {segment['end_hms']}",
            #     level=2
            # )

        # Insert time period description
        # time_para = doc.add_paragraph(f"Segment {segment['segment_index']} text: {segment['start_hms']} - {segment['end_hms']}")
        # time_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # for run in time_para.runs:
        #     run.font.name = "Arial Unicode MS"
        #     run.font.size = Pt(11)
        #     run.font.bold = True

        # Insert all text content for this time period (combined into one paragraph)
        if segment["text_segments"]:
            # Connect all text segments into one long text
            combined_text = " ".join([text_seg["text"] for text_seg in segment["text_segments"]])
            text_para = doc.add_paragraph(combined_text)
            text_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justified alignment
            text_para.paragraph_format.line_spacing = 1.2  # Line spacing 1.2x
            for run in text_para.runs:
                run.font.name = "Arial Unicode MS"
                run.font.size = Pt(12)
        # else:
        #     # If no text in this time period, add a note
        #     empty_para = doc.add_paragraph("(No speech content in this time period)")
        #     empty_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        #     for run in empty_para.runs:
        #         run.font.name = "Arial Unicode MS"
        #         run.font.size = Pt(10)
        #         run.font.italic = True

        # Insert page break (except for last segment)
        if segment["segment_index"] < len(segments_data):
            doc.add_page_break()

    # Save document
    doc.save(output_path)
    print(f"Word document generation completed: {output_path}")

# ===================== Main Function =====================
def main():
    # Parse command line arguments
    parser = argparse.ArgumentParser(description='Video to Document Tool - Alibaba Cloud Speech Recognition Version')
    parser.add_argument('--use_cache', action='store_true', help='Use cached data, skip API calls')
    parser.add_argument('--video', type=str, default=VIDEO_PATH, help='Input video path')
    parser.add_argument('--output', type=str, default=OUTPUT_WORD, help='Output Word document path')
    args = parser.parse_args()

    # Update configuration parameters
    video_path = args.video
    output_word = args.output

    try:
        # 1. Create temporary directory
        create_temp_dir()

        # 2. Extract audio (keep this step, but Alibaba Cloud version uses pre-uploaded audio link)
        audio_path = os.path.join(TEMP_DIR, "temp_audio.wav")
        video_to_audio(video_path, audio_path)

        # 3. Speech to text (using Alibaba Cloud API)
        cache_file_path = os.path.join(TEMP_DIR, "ali_api_result.json")
        if args.use_cache:
            print("Using cache mode: skipping API calls, using cached data directly")
        text_segments = ali_audio_to_text(audio_path, cache_file_path, use_cache=args.use_cache)

        # 4. Extract key frames (with deduplication)
        key_frames = extract_key_frames(video_path)

        # 4.5. Crop key frame black borders (executed after SSIM deduplication)
        crop_all_keyframes(key_frames)

        # 5. Match text and frames (simplified logic)
        segments_data = match_text_and_frames(text_segments, key_frames, video_path)

        # 6. Text segmentation (simplified logic)
        split_data = split_text_by_keyframes(segments_data)

        # 7. Generate Word document (simplified format)
        generate_word(split_data, output_word)

    except Exception as e:
        print(f"Execution error: {str(e)}")
    finally:
        # Clean up temporary files (optional, comment this line to keep temporary files for debugging)
        delete_temp_files()

        print("=== All processes completed! ===\n")


if __name__ == "__main__":
    main()