import os
import numpy as np
from moviepy.editor import VideoFileClip
import whisper
import cv2
from skimage.metrics import structural_similarity as ssim
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import timedelta
import shutil


# ===================== Configuration Parameters (Adjustable) =====================
VIDEO_PATH = "input_video.mp4"  # Input video path
OUTPUT_WORD = "output.docx"  # Output Word document path
TEMP_DIR = "temp_files"  # Temporary files directory (audio, key frames)
WHISPER_MODEL = "tiny"  # Whisper model (tiny/base/small/medium/large)
FRAME_INTERVAL = 1  # Frame extraction interval (seconds/frame)
SSIM_THRESHOLD = 0.93  # Key frame similarity threshold (≥0.93 considered duplicate)
MERGE_SEC = 2  # Adjacent text segment merge threshold (merge if interval < 2 seconds)

# ===================== Utility Functions =====================
def create_temp_dir():
    """Create temporary directory"""
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)
    return TEMP_DIR

def delete_temp_dir():
    """Delete temporary directory"""
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)

def seconds_to_hms(seconds):
    """Convert seconds to HH:MM:SS format"""
    return str(timedelta(seconds=round(seconds)))

def video_to_audio(video_path, audio_path):
    """Extract audio from video (WAV format)"""
    with VideoFileClip(video_path) as video:
        audio = video.audio
        audio.write_audiofile(audio_path, codec="pcm_s16le")  # WAV uncompressed format, better for Whisper recognition
    print(f"Audio extraction completed: {audio_path}")

def audio_to_text(audio_path):
    """Speech to text (with timestamps)"""
    model = whisper.load_model(WHISPER_MODEL)
    result = model.transcribe(audio_path, language="zh", word_timestamps=False)  # Disable word-level timestamps, keep segment-level
    segments = result["segments"]

    print(f"Speech recognition raw segments: {len(segments)} total")

    # Use original segments directly, no merging
    # Format output: add HH:MM:SS timestamps
    formatted_segments = []
    for seg in segments:
        formatted_segments.append({
            "start": seg["start"],
            "end": seg["end"],
            "start_hms": seconds_to_hms(seg["start"]),
            "end_hms": seconds_to_hms(seg["end"]),
            "text": seg["text"].strip()
        })
    print(f"Speech to text completed, {len(formatted_segments)} raw segments")
    print("Segment details:")
    for i, seg in enumerate(formatted_segments):
        print(f"  Segment {i+1}: {seg['start_hms']} - {seg['end_hms']}: {seg['text'][:50]}...")
    return formatted_segments

def extract_key_frames(video_path):
    """Extract key frames (with deduplication)"""
    cap = cv2.VideoCapture(video_path)
    fps = cap.get(cv2.CAP_PROP_FPS)  # Video frame rate
    frame_interval = int(fps * FRAME_INTERVAL)  # Extract one frame every FRAME_INTERVAL seconds (converted to frame count)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

    key_frames = []  # Store key frames: [{time, path}, ...]
    last_frame_gray = None  # Previous key frame grayscale image
    frame_count = 0

    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break

        # Extract frames at intervals
        if frame_count % frame_interval != 0:
            frame_count += 1
            continue

        # Calculate current frame timestamp (seconds)
        current_time = frame_count / fps

        # Convert to grayscale (for SSIM calculation)
        current_frame_gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)

        # Determine if it's a key frame (first frame is always kept)
        is_key_frame = False
        if last_frame_gray is None:
            is_key_frame = True
        else:
            # Calculate SSIM similarity (win_size must be odd, adapt to different frame sizes)
            win_size = min(7, current_frame_gray.shape[0]//2, current_frame_gray.shape[1]//2)
            if win_size % 2 == 0:
                win_size -= 1
            sim, _ = ssim(last_frame_gray, current_frame_gray, full=True, win_size=win_size)
            if sim < SSIM_THRESHOLD:
                is_key_frame = True

        if is_key_frame:
            # Save key frame
            frame_path = os.path.join(TEMP_DIR, f"key_frame_{len(key_frames)}.png")
            cv2.imwrite(frame_path, frame)
            key_frames.append({
                "time": current_time,
                "path": frame_path
            })
            last_frame_gray = current_frame_gray
            print(f"Saved key frame: {frame_path} (time: {seconds_to_hms(current_time)})")

        frame_count += 1

    cap.release()
    print(f"Key frame extraction completed, {len(key_frames)} frames")
    return key_frames

def get_video_duration(video_path):
    """Get video total duration (seconds)"""
    with VideoFileClip(video_path) as video:
        duration = video.duration
    return duration

def match_text_and_frames(text_segments, key_frames, video_path):
    """Segment text fragments using key frame times: simplified logic"""
    # Get video total duration
    video_duration = get_video_duration(video_path)

    # Extract key frame times and sort them
    key_frame_times = sorted([kf["time"] for kf in key_frames])

    print(f"Key frame segmentation started: {len(text_segments)} text segments, {len(key_frames)} key frames")
    print(f"Video total duration: {seconds_to_hms(video_duration)}")
    print(f"Key frame times: {key_frame_times}")

    # Create segment data: each key frame corresponds to one segment
    segments_data = []

    # If no key frames, create one segment containing all text fragments
    if not key_frame_times:
        segments_data.append({
            "segment_index": 1,
            "start_time": 0,
            "end_time": video_duration,
            "start_hms": seconds_to_hms(0),
            "end_hms": seconds_to_hms(video_duration),
            "key_frame": None,
            "text_segments": text_segments
        })
        print(f"  Segment 1: {seconds_to_hms(0)} - {seconds_to_hms(video_duration)} -> {len(text_segments)} text segments [No key frame]")
    else:
        # First segment: from 0 to second key frame time (if second key frame exists)
        first_keyframe = key_frames[0]
        end_time = key_frame_times[1] if len(key_frame_times) > 1 else video_duration

        # Find all text fragments belonging to this time period
        segment_texts = []
        for seg in text_segments:
            if 0 <= seg["start"] < end_time:
                segment_texts.append(seg)

        segments_data.append({
            "segment_index": 1,
            "start_time": 0,
            "end_time": end_time,
            "start_hms": seconds_to_hms(0),
            "end_hms": seconds_to_hms(end_time),
            "key_frame": first_keyframe,
            "text_segments": segment_texts
        })
        print(f"  Segment 1: {seconds_to_hms(0)} - {seconds_to_hms(end_time)} -> {len(segment_texts)} text segments [Key frame: {first_keyframe['time']:.1f}s]")

        # Middle segments: from i-th key frame to (i+1)-th key frame (i from 1 to n-2)
        for i in range(1, len(key_frame_times) - 1):
            start_time = key_frame_times[i]
            end_time = key_frame_times[i + 1]
            key_frame = key_frames[i]

            # Find all text fragments belonging to this time period
            segment_texts = []
            for seg in text_segments:
                if start_time <= seg["start"] < end_time:
                    segment_texts.append(seg)

            segments_data.append({
                "segment_index": i + 1,
                "start_time": start_time,
                "end_time": end_time,
                "start_hms": seconds_to_hms(start_time),
                "end_hms": seconds_to_hms(end_time),
                "key_frame": key_frame,
                "text_segments": segment_texts
            })
            print(f"  Segment {i+1}: {seconds_to_hms(start_time)} - {seconds_to_hms(end_time)} -> {len(segment_texts)} text segments [Key frame: {key_frame['time']:.1f}s]")

        # Last segment: from last key frame to video end
        if len(key_frame_times) > 1:
            last_keyframe = key_frames[-1]
            start_time = key_frame_times[-1]
            end_time = video_duration

            # Find all text fragments belonging to this time period
            segment_texts = []
            for seg in text_segments:
                if start_time <= seg["start"] <= end_time:
                    segment_texts.append(seg)

            segments_data.append({
                "segment_index": len(key_frame_times),
                "start_time": start_time,
                "end_time": end_time,
                "start_hms": seconds_to_hms(start_time),
                "end_hms": seconds_to_hms(end_time),
                "key_frame": last_keyframe,
                "text_segments": segment_texts
            })
            print(f"  Segment {len(key_frame_times)}: {seconds_to_hms(start_time)} - {seconds_to_hms(end_time)} -> {len(segment_texts)} text segments [Key frame: {last_keyframe['time']:.1f}s]")

    # Validation: ensure all text fragments are assigned to some segment
    assigned_texts = sum(len(seg["text_segments"]) for seg in segments_data)
    if assigned_texts != len(text_segments):
        raise Exception(f"Text fragment assignment error: {len(text_segments)} total fragments, only {assigned_texts} assigned")

    print(f"Key frame segmentation completed, {len(segments_data)} segments")
    return segments_data

def split_text_by_keyframes(segments_data):
    """Simplify text segmentation: directly use segment data"""
    # In the new segmentation logic, segmentation is already done, directly return segment data
    print(f"Text segmentation completed: {len(segments_data)} segments")
    return segments_data

def generate_word(segments_data, output_path):
    """Generate Word document: simplified format - key frames + corresponding text segments"""
    doc = Document()

    print(f"Generating Word document: {len(segments_data)} segments")

    # Set document title
    title = doc.add_heading("Educational Video Notes", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Traverse segment data and insert content
    for segment in segments_data:
        # Debug information
        print(f"Generating segment {segment['segment_index']}: key_frame={segment['key_frame'] is not None}")
        if segment["key_frame"]:
            print(f"  Key frame time: {segment['key_frame']['time']:.1f}s")

        # Insert key frame title
        if segment["key_frame"]:
            keyframe_heading = doc.add_heading(
                f"Key Frame {segment['segment_index']} ({segment['key_frame']['time']:.1f}s)",
                level=2
            )

            # Insert key frame image
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(segment["key_frame"]["path"], width=Inches(5))
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            # For segments without key frames (like the last segment), only show time period
            time_heading = doc.add_heading(
                f"Segment {segment['segment_index']}: {segment['start_hms']} - {segment['end_hms']}",
                level=2
            )

        # Insert time period description
        time_para = doc.add_paragraph(f"Segment {segment['segment_index']} text: {segment['start_hms']} - {segment['end_hms']}")
        time_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in time_para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            run.font.bold = True

        # Insert all text content for this time period
        if segment["text_segments"]:
            for text_seg in segment["text_segments"]:
                text_para = doc.add_paragraph(text_seg["text"])
                text_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Justified alignment
                text_para.paragraph_format.line_spacing = 1.5  # Line spacing 1.5x
                for run in text_para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
        else:
            # If no text in this time period, add a note
            empty_para = doc.add_paragraph("(No speech content in this time period)")
            empty_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in empty_para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.font.italic = True

        # Insert page break (except for last segment)
        if segment["segment_index"] < len(segments_data):
            doc.add_page_break()

    # Save document
    doc.save(output_path)
    print(f"Word document generation completed: {output_path}")

# ===================== Main Function =====================
def main():
    try:
        # 1. Create temporary directory
        create_temp_dir()

        # 2. Extract audio
        audio_path = os.path.join(TEMP_DIR, "temp_audio.wav")
        video_to_audio(VIDEO_PATH, audio_path)

        # 3. Speech to text (with timestamps)
        text_segments = audio_to_text(audio_path)

        # 4. Extract key frames (with deduplication)
        key_frames = extract_key_frames(VIDEO_PATH)

        # 5. Match text and frames (simplified logic)
        segments_data = match_text_and_frames(text_segments, key_frames, VIDEO_PATH)

        # 6. Text segmentation (simplified logic)
        split_data = split_text_by_keyframes(segments_data)

        # 7. Generate Word document (simplified format)
        generate_word(split_data, OUTPUT_WORD)

        print("\n=== All processes completed! ===")
    except Exception as e:
        print(f"Execution error: {str(e)}")
    finally:
        # Clean up temporary files (optional, comment this line to keep temporary files for debugging)
        delete_temp_dir()

if __name__ == "__main__":
    main()
